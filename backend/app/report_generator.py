"""
보고서 생성기 모듈

분석 결과를 Word/PDF 보고서로 변환합니다.

주요 기능:
- 이슈 → 액션 카드 구조화
- Word (docx) 보고서 생성
- PDF 보고서 생성 (docx 변환)
- 템플릿 기반 커스터마이징
"""

from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
import json
import re

# python-docx 가용성 체크
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[Report] python-docx가 설치되지 않았습니다. pip install python-docx")


# ============================================================
# 이슈 → 액션 카드 타입 정의
# ============================================================

@dataclass
class Evidence:
    """근거 자료"""
    document: str           # 문서명
    article: str            # 조항
    page: Optional[str]     # 페이지
    url: Optional[str]      # 원본 URL


@dataclass
class Impact:
    """영향도"""
    cost: str               # 비용 영향 (높음/중간/낮음 또는 금액)
    operation: str          # 운영 영향
    admin: str              # 행정 부담
    reputation: str         # 평판 리스크


@dataclass
class ActionItem:
    """실행 항목"""
    category: str           # immediate/shortTerm/midTerm/longTerm
    description: str
    deadline: Optional[str]
    priority: str           # high/medium/low
    responsible: Optional[str]


@dataclass
class IssueActionCard:
    """이슈 액션 카드"""
    id: str
    title: str
    
    # 핵심 정보
    obligation: str         # 무엇을 해야 하는가
    deadline: str           # 언제까지
    
    # 근거
    evidences: List[Evidence]
    
    # 영향도
    impact: Impact
    
    # 실행 계획
    actions: List[ActionItem]
    
    # 담당
    owner: str              # 담당부서
    
    # 메타
    importance_score: float
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())


# ============================================================
# 이슈 → 액션 카드 변환
# ============================================================

def convert_issue_to_action_card(
    issue: Dict[str, Any],
    analysis_result: Optional[Dict[str, Any]] = None
) -> IssueActionCard:
    """
    이슈 데이터를 액션 카드로 변환
    
    Args:
        issue: 발굴된 이슈 데이터
        analysis_result: 심층 분석 결과 (선택)
    
    Returns:
        IssueActionCard
    """
    # 기본 정보 추출
    title = issue.get("title", "")
    summary = issue.get("summary", "")
    score = issue.get("score", {})
    sources = issue.get("sources", [])
    
    # 근거 자료 변환
    evidences = []
    for src in sources[:5]:
        evidences.append(Evidence(
            document=src.get("docTitle", src.get("orgName", "")),
            article="",
            page=None,
            url=None,
        ))
    
    # 영향도 추정 (점수 기반)
    total_score = score.get("total", 0.5)
    impact = Impact(
        cost="높음" if total_score > 0.7 else "중간" if total_score > 0.4 else "낮음",
        operation="검토 필요",
        admin="검토 필요",
        reputation="검토 필요",
    )
    
    # 분석 결과에서 추가 정보 추출
    obligation = summary
    deadline = "확인 필요"
    owner = "환경안전팀"
    actions = []
    
    if analysis_result:
        # 분석 결과에서 의무사항, 마감일 등 추출
        for step in analysis_result.get("steps", []):
            content = step.get("content", "")
            
            # 마감일 추출 시도
            date_match = re.search(r'(\d{4}년\s*\d{1,2}월|\d{4}\.\d{1,2}\.\d{1,2})', content)
            if date_match and deadline == "확인 필요":
                deadline = date_match.group(1)
            
            # 실행 항목 추출 (단계별)
            if step.get("stepName") == "결론 도출" or "대응" in step.get("stepName", ""):
                # 번호 목록 형태의 권고사항 추출
                action_matches = re.findall(r'\d+\.\s+([^\n]+)', content)
                for i, action_desc in enumerate(action_matches[:5]):
                    category = "immediate" if i < 2 else "shortTerm" if i < 4 else "midTerm"
                    actions.append(ActionItem(
                        category=category,
                        description=action_desc.strip(),
                        deadline=None,
                        priority="high" if i < 2 else "medium",
                        responsible=owner,
                    ))
    
    # 기본 액션 항목 (분석 결과 없을 때)
    if not actions:
        actions = [
            ActionItem(
                category="immediate",
                description="관련 규정 확인 및 현황 파악",
                deadline=None,
                priority="high",
                responsible=owner,
            ),
            ActionItem(
                category="shortTerm",
                description="내부 담당자 공유 및 검토",
                deadline=None,
                priority="medium",
                responsible=owner,
            ),
        ]
    
    return IssueActionCard(
        id=issue.get("id", ""),
        title=title,
        obligation=obligation,
        deadline=deadline,
        evidences=evidences,
        impact=impact,
        actions=actions,
        owner=owner,
        importance_score=total_score,
    )


# ============================================================
# 보고서 생성
# ============================================================

class ReportGenerator:
    """보고서 생성기"""
    
    def __init__(self, template_path: Optional[str] = None):
        """
        Args:
            template_path: 템플릿 docx 파일 경로 (선택)
        """
        self.template_path = template_path
    
    def generate_word_report(
        self,
        title: str,
        action_cards: List[IssueActionCard],
        metadata: Dict[str, Any] = None,
        output_path: str = None
    ) -> Optional[str]:
        """
        Word 보고서 생성
        
        Args:
            title: 보고서 제목
            action_cards: 이슈 액션 카드 목록
            metadata: 보고서 메타데이터
            output_path: 출력 경로 (None이면 자동 생성)
        
        Returns:
            생성된 파일 경로
        """
        if not DOCX_AVAILABLE:
            print("[Report] python-docx가 설치되지 않아 보고서를 생성할 수 없습니다.")
            return None
        
        # 문서 생성
        if self.template_path and Path(self.template_path).exists():
            doc = Document(self.template_path)
        else:
            doc = Document()
        
        # 스타일 설정
        self._setup_styles(doc)
        
        # 표지
        self._add_cover_page(doc, title, metadata)
        
        # 요약
        self._add_executive_summary(doc, action_cards)
        
        # 이슈별 상세
        for i, card in enumerate(action_cards, 1):
            self._add_issue_section(doc, card, i)
        
        # 부록
        self._add_appendix(doc, action_cards)
        
        # 저장
        if not output_path:
            output_dir = Path(__file__).parent.parent.parent / "frontend" / "data" / "reports"
            output_dir.mkdir(parents=True, exist_ok=True)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = str(output_dir / f"report_{timestamp}.docx")
        
        doc.save(output_path)
        return output_path
    
    def _setup_styles(self, doc: Document):
        """문서 스타일 설정"""
        styles = doc.styles
        
        # 제목 스타일
        try:
            title_style = styles.add_style('ReportTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.paragraph_format.space_after = Pt(12)
        except:
            pass
        
        # 소제목 스타일
        try:
            heading_style = styles.add_style('IssueHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True
            heading_style.paragraph_format.space_before = Pt(18)
            heading_style.paragraph_format.space_after = Pt(6)
        except:
            pass
    
    def _add_cover_page(self, doc: Document, title: str, metadata: Dict[str, Any] = None):
        """표지 추가"""
        # 제목
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(28)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # 메타데이터
        if metadata:
            meta_p = doc.add_paragraph()
            meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if metadata.get("date"):
                meta_p.add_run(f"작성일: {metadata['date']}\n")
            if metadata.get("author"):
                meta_p.add_run(f"작성자: {metadata['author']}\n")
            if metadata.get("department"):
                meta_p.add_run(f"부서: {metadata['department']}")
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, action_cards: List[IssueActionCard]):
        """요약 섹션 추가"""
        doc.add_heading("요약", level=1)
        
        # 전체 통계
        total = len(action_cards)
        high_priority = sum(1 for c in action_cards if c.importance_score > 0.7)
        
        summary_text = f"""
본 보고서는 총 {total}건의 주요 이슈를 분석한 결과입니다.
- 높은 중요도: {high_priority}건
- 중간 중요도: {total - high_priority}건

각 이슈에 대한 의무사항, 영향도, 실행 계획을 상세히 기술하였습니다.
"""
        doc.add_paragraph(summary_text.strip())
        
        # 이슈 목록 테이블
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # 헤더
        header_cells = table.rows[0].cells
        header_cells[0].text = "No."
        header_cells[1].text = "이슈명"
        header_cells[2].text = "중요도"
        header_cells[3].text = "마감일"
        
        # 데이터
        for i, card in enumerate(action_cards, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = card.title[:50]
            row_cells[2].text = f"{card.importance_score:.0%}"
            row_cells[3].text = card.deadline
        
        doc.add_page_break()
    
    def _add_issue_section(self, doc: Document, card: IssueActionCard, number: int):
        """이슈 섹션 추가"""
        # 이슈 제목
        doc.add_heading(f"{number}. {card.title}", level=1)
        
        # 핵심 정보
        doc.add_heading("핵심 정보", level=2)
        
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'
        
        info_table.rows[0].cells[0].text = "의무사항"
        info_table.rows[0].cells[1].text = card.obligation[:200] if card.obligation else "-"
        
        info_table.rows[1].cells[0].text = "마감일"
        info_table.rows[1].cells[1].text = card.deadline
        
        info_table.rows[2].cells[0].text = "담당부서"
        info_table.rows[2].cells[1].text = card.owner
        
        # 영향도
        doc.add_heading("영향도 분석", level=2)
        
        impact_table = doc.add_table(rows=4, cols=2)
        impact_table.style = 'Table Grid'
        
        impact_table.rows[0].cells[0].text = "비용 영향"
        impact_table.rows[0].cells[1].text = card.impact.cost
        
        impact_table.rows[1].cells[0].text = "운영 영향"
        impact_table.rows[1].cells[1].text = card.impact.operation
        
        impact_table.rows[2].cells[0].text = "행정 부담"
        impact_table.rows[2].cells[1].text = card.impact.admin
        
        impact_table.rows[3].cells[0].text = "평판 리스크"
        impact_table.rows[3].cells[1].text = card.impact.reputation
        
        # 실행 계획
        doc.add_heading("실행 계획", level=2)
        
        for action in card.actions:
            category_label = {
                "immediate": "즉시",
                "shortTerm": "단기",
                "midTerm": "중기",
                "longTerm": "장기",
            }.get(action.category, action.category)
            
            priority_label = {
                "high": "🔴",
                "medium": "🟡",
                "low": "🟢",
            }.get(action.priority, "")
            
            p = doc.add_paragraph()
            p.add_run(f"[{category_label}] {priority_label} ").bold = True
            p.add_run(action.description)
        
        # 근거 자료
        if card.evidences:
            doc.add_heading("근거 자료", level=2)
            for evidence in card.evidences[:5]:
                doc.add_paragraph(f"• {evidence.document}", style='List Bullet')
        
        doc.add_page_break()
    
    def _add_appendix(self, doc: Document, action_cards: List[IssueActionCard]):
        """부록 추가"""
        doc.add_heading("부록", level=1)
        
        # 용어 설명
        doc.add_heading("용어 설명", level=2)
        doc.add_paragraph("• 중요도: 법적 강제성, 신규성, 파급력, 국제 동향을 종합한 점수")
        doc.add_paragraph("• 액션 카드: 이슈별 실행 계획을 구조화한 문서")
        
        # 작성 정보
        doc.add_heading("작성 정보", level=2)
        doc.add_paragraph(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("생성 도구: EcoMonitor AI RAG 분석 시스템")


# ============================================================
# 헬퍼 함수
# ============================================================

def generate_report_from_session(
    session_data: Dict[str, Any],
    analysis_results: Optional[Dict[str, Any]] = None,
    output_format: str = "docx"
) -> Optional[str]:
    """
    세션 데이터에서 보고서 생성
    
    Args:
        session_data: 발굴 세션 데이터
        analysis_results: 심층 분석 결과
        output_format: 출력 형식 (docx, pdf)
    
    Returns:
        생성된 파일 경로
    """
    issues = session_data.get("issues", [])
    
    if not issues:
        return None
    
    # 이슈 → 액션 카드 변환
    action_cards = []
    for issue in issues:
        issue_id = issue.get("id", "")
        issue_analysis = None
        
        if analysis_results:
            issue_analysis = analysis_results.get(issue_id, {})
        
        card = convert_issue_to_action_card(issue, issue_analysis)
        action_cards.append(card)
    
    # 중요도순 정렬
    action_cards.sort(key=lambda x: x.importance_score, reverse=True)
    
    # 보고서 생성
    generator = ReportGenerator()
    
    metadata = {
        "date": datetime.now().strftime("%Y년 %m월 %d일"),
        "author": "EcoMonitor AI",
        "department": "환경규제분석팀",
    }
    
    session_name = session_data.get("name", "이슈 분석")
    title = f"{session_name} 보고서"
    
    return generator.generate_word_report(
        title=title,
        action_cards=action_cards,
        metadata=metadata,
    )
