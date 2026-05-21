"""
DOCX 텍스트 추출 모듈
- python-docx 기반 문서 파싱
- 문단, 표, 머리글/바닥글 추출
"""
import re
import time
from pathlib import Path
from typing import Optional, List, Tuple

from .base import (
    BaseExtractor,
    ExtractionResult,
    ExtractionStatus,
    ExtractionMethod,
    PageResult,
    QualityDetails,
    TableData,
    detect_utf16_misencoding,
    quick_validate_text_quality,
    recover_encoding_with_retry
)
from ..table_recovery.table_builder import TableBuilder, Table, TableCell
from ..config import ExtractionConfig, default_config
from typing import Dict

# python-docx 지연 로딩
_docx = None


def get_docx():
    """python-docx 모듈 지연 로딩"""
    global _docx
    if _docx is None:
        try:
            import docx
            _docx = docx
        except ImportError:
            print("[WARNING] python-docx not installed. DOCX extraction will fail.")
            return None
    return _docx


class DOCXExtractor(BaseExtractor):
    """
    DOCX 텍스트 추출기
    
    DOCX 파일 구조:
    - ZIP 압축 파일 (OOXML 형식)
    - word/document.xml에 본문
    - word/header*.xml, word/footer*.xml에 머리글/바닥글
    - word/styles.xml에 스타일 정보
    """
    
    SUPPORTED_FORMATS = ["docx", "doc"]
    
    KOREAN_PATTERN = re.compile(r'[가-힣]')
    BROKEN_CHAR_PATTERN = re.compile(r'[\ufffd\x00-\x08\x0b\x0c\x0e-\x1f]')
    
    def __init__(self, config: Optional[ExtractionConfig] = None):
        self.config = config or default_config
    
    def supports(self, file_format: str) -> bool:
        """DOCX 형식 지원 여부"""
        return file_format.lower() in self.SUPPORTED_FORMATS
    
    def extract(self, file_path: str) -> ExtractionResult:
        """
        DOCX 파일에서 텍스트 추출
        """
        start_time = time.time()
        file_path = Path(file_path)
        file_format = file_path.suffix.lower().lstrip('.')
        
        if not file_path.exists():
            return ExtractionResult(
                file_path=str(file_path),
                file_format=file_format,
                status=ExtractionStatus.FAILED,
                error_message=f"File not found: {file_path}"
            )
        
        # .doc 파일은 지원하지 않음 (바이너리 형식)
        if file_format == "doc":
            return ExtractionResult(
                file_path=str(file_path),
                file_format=file_format,
                status=ExtractionStatus.FAILED,
                error_message="Legacy .doc format is not supported. Please convert to .docx"
            )
        
        docx_module = get_docx()
        if docx_module is None:
            return ExtractionResult(
                file_path=str(file_path),
                file_format=file_format,
                status=ExtractionStatus.FAILED,
                error_message="python-docx module not available"
            )
        
        try:
            # 문서 열기
            doc = docx_module.Document(str(file_path))
            
            # 텍스트 추출
            texts = []
            extracted_tables = []
            has_tables = False
            
            # 1. 본문 문단 추출
            paragraphs = self._extract_paragraphs(doc)
            if paragraphs:
                texts.append(paragraphs)
            
            # 2. 표 추출 (마크다운 형식)
            if self.config.docx.extract_tables:
                extracted_tables = self._extract_tables_with_structure(doc)
                if extracted_tables:
                    has_tables = True
                    # 텍스트와 표 병합
                    combined_text = self._merge_text_and_tables(
                        "\n".join(texts) if texts else "",
                        extracted_tables
                    )
                    texts = [combined_text]
            else:
                # 기존 방식으로 표 추출
                table_text = self._extract_tables(doc)
                if table_text:
                    texts.append(table_text)
            
            # 3. 머리글/바닥글 추출 (선택적)
            headers_footers = self._extract_headers_footers(doc)
            if headers_footers:
                texts.append(headers_footers)
            
            text = "\n\n".join(texts)
            
            if not text or len(text.strip()) < 5:
                return ExtractionResult(
                    file_path=str(file_path),
                    file_format=file_format,
                    status=ExtractionStatus.FAILED,
                    error_message="No text content found in DOCX file"
                )
            
            # 후처리
            text, corrections = self._apply_postprocessing(text)
            
            # 인코딩 품질 검증 및 자동 복구
            is_valid, quality_score, error_reason = quick_validate_text_quality(text)
            
            if not is_valid or quality_score < 0.5:
                # 인코딩 오류 감지 - 복구 시도
                print(f"[DOCX] 인코딩 품질 낮음 ({quality_score:.2f}): {error_reason}")
                
                recovered_text, recovery_corrections = recover_encoding_with_retry(text)
                _, new_score, _ = quick_validate_text_quality(recovered_text)
                
                if new_score > quality_score:
                    text = recovered_text
                    corrections.extend(recovery_corrections)
                    print(f"[DOCX] 인코딩 복구 성공: {quality_score:.2f} -> {new_score:.2f}")
            
            # 품질 검증
            quality_details = self._validate_quality(text)
            
            # 토큰 수 계산
            token_count = self._count_tokens(text)
            
            processing_time_ms = int((time.time() - start_time) * 1000)
            
            # 상태 결정
            if quality_details.overall_score >= self.config.quality_validation.pass_threshold:
                status = ExtractionStatus.COMPLETED
            elif quality_details.overall_score >= self.config.quality_validation.llm_fallback_threshold:
                status = ExtractionStatus.COMPLETED
            else:
                status = ExtractionStatus.LLM_FALLBACK if self.config.quality_validation.auto_llm_fallback else ExtractionStatus.FAILED
            
            return ExtractionResult(
                file_path=str(file_path),
                file_format=file_format,
                status=status,
                text=text,
                pages=[PageResult(
                    page_num=1, 
                    text=text, 
                    method=ExtractionMethod.TEXT_LAYER,
                    has_tables=has_tables,
                    tables=extracted_tables
                )],
                page_count=1,
                quality_score=quality_details.overall_score,
                quality_details=quality_details,
                extraction_method=ExtractionMethod.TEXT_LAYER,
                char_count=len(text),
                token_count=token_count,
                processing_time_ms=processing_time_ms,
                corrections=corrections
            )
            
        except Exception as e:
            return ExtractionResult(
                file_path=str(file_path),
                file_format=file_format,
                status=ExtractionStatus.FAILED,
                error_message=str(e),
                processing_time_ms=int((time.time() - start_time) * 1000)
            )
    
    def _extract_paragraphs(self, doc) -> str:
        """본문 문단 추출"""
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 스타일 기반 구조화 (제목 레벨 등)
                style_name = para.style.name if para.style else ""
                
                if "Heading" in style_name or "제목" in style_name:
                    # 제목은 앞뒤로 여백 추가
                    paragraphs.append(f"\n{text}\n")
                else:
                    paragraphs.append(text)
        
        return "\n".join(paragraphs)
    
    def _extract_tables(self, doc) -> str:
        """표 추출 (기존 방식 - 호환성 유지)"""
        table_texts = []
        
        for table_idx, table in enumerate(doc.tables):
            rows = []
            
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        cells.append(cell_text)
                
                if cells:
                    rows.append(" | ".join(cells))
            
            if rows:
                table_text = f"[표 {table_idx + 1}]\n" + "\n".join(rows)
                table_texts.append(table_text)
        
        return "\n\n".join(table_texts)
    
    def _get_cell_span_info(self, cell) -> Tuple[int, int]:
        """
        DOCX 셀의 gridSpan, vMerge 속성 파싱
        
        Args:
            cell: python-docx Cell 객체
            
        Returns:
            (col_span, row_span)
        """
        col_span = 1
        row_span = 1
        
        try:
            tc = cell._tc
            tcPr = tc.tcPr
            
            if tcPr is not None:
                # 열 병합 (gridSpan)
                gridSpan = tcPr.gridSpan
                if gridSpan is not None:
                    col_span = gridSpan.val
                
                # 행 병합 (vMerge) - 시작 셀인 경우만 처리
                # vMerge가 있고 val이 'restart'이면 병합 시작
                # vMerge가 있고 val이 없거나 'continue'이면 병합 계속
                vMerge = tcPr.vMerge
                if vMerge is not None:
                    # restart: 병합 시작, None/continue: 이전 셀의 연속
                    if vMerge.val == 'restart' or vMerge.val is None:
                        row_span = 1  # 실제 행 병합 수는 후처리에서 계산
        except Exception:
            pass
        
        return col_span, row_span
    
    def _is_cell_merge_continue(self, cell) -> bool:
        """셀이 행 병합의 연속인지 확인 (vMerge continue)"""
        try:
            tc = cell._tc
            tcPr = tc.tcPr
            
            if tcPr is not None:
                vMerge = tcPr.vMerge
                if vMerge is not None:
                    # val이 없거나 'continue'이면 병합 연속
                    return vMerge.val is None or vMerge.val == 'continue'
        except Exception:
            pass
        
        return False
    
    def _expand_merged_cells(self, cells_with_span: List[List[Dict]]) -> List[List[str]]:
        """
        병합 셀을 확장하여 정규 2D 배열로 변환
        
        Args:
            cells_with_span: [
                [{"text": "헤더1", "col_span": 2, "row_span": 1}],
                [{"text": "값1", "col_span": 1, "row_span": 1},
                 {"text": "값2", "col_span": 1, "row_span": 1}]
            ]
            
        Returns:
            [["헤더1", ""], ["값1", "값2"]]
        """
        if not cells_with_span:
            return []
        
        # 최대 열/행 수 계산
        max_cols = 0
        for row in cells_with_span:
            row_cols = sum(cell.get('col_span', 1) for cell in row)
            max_cols = max(max_cols, row_cols)
        
        max_rows = len(cells_with_span)
        for row_idx, row in enumerate(cells_with_span):
            for cell in row:
                row_span = cell.get('row_span', 1)
                if row_idx + row_span > max_rows:
                    max_rows = row_idx + row_span
        
        # 빈 그리드 생성
        grid = [[None for _ in range(max_cols)] for _ in range(max_rows)]
        
        # 셀 배치
        for row_idx, row in enumerate(cells_with_span):
            col_idx = 0
            for cell in row:
                # 이미 채워진 셀 건너뛰기
                while col_idx < max_cols and grid[row_idx][col_idx] is not None:
                    col_idx += 1
                if col_idx >= max_cols:
                    break
                
                text = cell.get('text', '')
                col_span = cell.get('col_span', 1)
                row_span = cell.get('row_span', 1)
                is_merge_continue = cell.get('is_merge_continue', False)
                
                # 병합 연속 셀은 빈 문자열로 처리
                if is_merge_continue:
                    grid[row_idx][col_idx] = ''
                    col_idx += 1
                    continue
                
                # 첫 셀에 텍스트 배치
                grid[row_idx][col_idx] = text
                
                # 병합 영역에 빈 문자열 배치
                for r in range(row_span):
                    for c in range(col_span):
                        target_row = row_idx + r
                        target_col = col_idx + c
                        if target_row < max_rows and target_col < max_cols:
                            if r == 0 and c == 0:
                                continue  # 첫 셀은 이미 설정됨
                            grid[target_row][target_col] = ''
                
                col_idx += col_span
        
        # None을 빈 문자열로 변환
        return [[cell if cell is not None else '' for cell in row] for row in grid]
    
    def _is_valid_table(self, table: List[List[str]]) -> bool:
        """
        유효한 표인지 검사 (노이즈 필터링)
        
        검증 기준:
        1. 최소 행 수 (config.docx.min_table_rows)
        2. 최소 열 수 (config.docx.min_table_cols)
        3. 빈 셀 비율 30% 이상 채워짐
        """
        if not table:
            return False
        
        row_count = len(table)
        col_count = max(len(row) for row in table) if table else 0
        
        if row_count < self.config.docx.min_table_rows:
            return False
        if col_count < self.config.docx.min_table_cols:
            return False
        
        # 빈 셀이 너무 많은 표 제외
        total_cells = row_count * col_count
        non_empty_cells = sum(
            1 for row in table for cell in row 
            if cell and str(cell).strip()
        )
        
        if total_cells > 0 and non_empty_cells / total_cells < 0.3:
            return False
        
        return True
    
    def _clean_cell(self, cell) -> str:
        """셀 데이터 정리"""
        if cell is None:
            return ''
        text = str(cell).strip()
        text = text.replace('\n', ' ')      # 줄바꿈 제거
        text = text.replace('|', '\\|')     # 파이프 이스케이프
        text = re.sub(r'\s+', ' ', text)    # 연속 공백 정리
        return text
    
    def _table_to_markdown(self, table: List[List[str]]) -> str:
        """
        2D 배열을 마크다운 테이블로 변환
        
        특징:
        - 열 수 자동 정규화 (불균일 행 처리)
        - 셀 내 줄바꿈 → 공백 변환
        - 파이프(|) 문자 이스케이프
        """
        if not table or not table[0]:
            return ""
        
        lines = []
        col_count = max(len(row) for row in table)
        
        # 헤더 행
        header = table[0]
        header_cells = [self._clean_cell(c) for c in header]
        header_cells += [''] * (col_count - len(header_cells))
        lines.append('| ' + ' | '.join(header_cells) + ' |')
        
        # 구분선
        lines.append('|' + '|'.join(['---'] * col_count) + '|')
        
        # 데이터 행
        for row in table[1:]:
            cells = [self._clean_cell(c) for c in row]
            cells += [''] * (col_count - len(cells))
            lines.append('| ' + ' | '.join(cells) + ' |')
        
        return '\n'.join(lines)
    
    def _extract_tables_with_structure(self, doc) -> List[TableData]:
        """
        표 구조 + 병합 정보 추출 (RAG 최적화)
        
        Args:
            doc: python-docx Document 객체
            
        Returns:
            List[TableData]: 추출된 표 목록
        """
        tables = []
        table_builder = TableBuilder()
        
        if not self.config.docx.extract_tables:
            return tables
        
        for table_idx, table in enumerate(doc.tables):
            try:
                cells_with_span = []
                
                for row in table.rows:
                    row_cells_with_span = []
                    
                    # 중복 셀 제거를 위한 처리 (python-docx는 병합 셀에서 같은 객체 반환)
                    seen_cells = set()
                    
                    for cell in row.cells:
                        # 동일한 셀 객체 중복 방지
                        cell_id = id(cell._tc)
                        if cell_id in seen_cells:
                            continue
                        seen_cells.add(cell_id)
                        
                        cell_text = cell.text.strip()
                        col_span, row_span = self._get_cell_span_info(cell)
                        is_merge_continue = self._is_cell_merge_continue(cell)
                        
                        row_cells_with_span.append({
                            'text': cell_text if not is_merge_continue else '',
                            'col_span': col_span,
                            'row_span': row_span,
                            'is_merge_continue': is_merge_continue
                        })
                    
                    if row_cells_with_span:
                        cells_with_span.append(row_cells_with_span)
                
                # 병합 셀 확장
                rows_data = self._expand_merged_cells(cells_with_span)
                
                # 표 유효성 검사 및 변환
                if rows_data and self._is_valid_table(rows_data):
                    # TableBuilder를 사용하여 Table 객체 생성
                    table_obj = table_builder.build_from_grid(rows_data, header_rows=1)
                    
                    tables.append(TableData(
                        table_index=table_idx,
                        rows=rows_data,
                        semantic_text=table_obj.to_semantic_text(),
                        structured_data=table_obj.to_structured_dict(),
                        markdown=table_obj.to_markdown(),  # 호환성용
                        row_count=table_obj.rows,
                        col_count=table_obj.cols,
                        page_num=1  # DOCX는 페이지 경계가 명확하지 않아 기본값 사용
                    ))
            
            except Exception as e:
                print(f"[DOCX] Table extraction error for table {table_idx}: {e}")
        
        return tables
    
    def _merge_text_and_tables(self, text: str, tables: List[TableData]) -> str:
        """
        텍스트와 표를 병합 (RAG 최적화된 semantic_text 우선 사용)
        
        현재 전략: 표 semantic_text를 텍스트 끝에 순차 추가
        
        Args:
            text: 본문 텍스트
            tables: 추출된 표 목록
            
        Returns:
            병합된 텍스트
        """
        if not tables:
            return text
        
        result_parts = [text.strip()]
        table_builder = TableBuilder()
        
        for table in tables:
            result_parts.append(f"\n\n<!-- 표 {table.table_index + 1} -->\n")
            
            # semantic_text 우선 사용 (RAG 최적화)
            if table.semantic_text:
                result_parts.append(table.semantic_text)
            elif table.markdown:
                result_parts.append(table.markdown)
            else:
                # 최종 폴백: 2D 그리드에서 직접 생성
                table_obj = table_builder.build_from_grid(table.rows, header_rows=1)
                result_parts.append(table_obj.to_semantic_text())
        
        return '\n'.join(result_parts)
    
    def _extract_headers_footers(self, doc) -> str:
        """머리글/바닥글 추출"""
        texts = []
        
        try:
            # 섹션별 머리글/바닥글 추출
            for section in doc.sections:
                # 머리글
                header = section.header
                if header and header.paragraphs:
                    header_text = " ".join(p.text.strip() for p in header.paragraphs if p.text.strip())
                    if header_text:
                        texts.append(f"[머리글] {header_text}")
                
                # 바닥글
                footer = section.footer
                if footer and footer.paragraphs:
                    footer_text = " ".join(p.text.strip() for p in footer.paragraphs if p.text.strip())
                    if footer_text:
                        texts.append(f"[바닥글] {footer_text}")
        except Exception:
            pass
        
        return "\n".join(texts)
    
    def _apply_postprocessing(self, text: str) -> Tuple[str, List[dict]]:
        """후처리 적용"""
        corrections = []
        
        # 연속 줄바꿈 정리
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # 연속 공백 정리
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\t+', ' ', text)
        
        # 특수 문자 정리
        special_chars = [
            (r'\u200b', ''),  # Zero-width space
            (r'\u00a0', ' '),  # Non-breaking space
            (r'\ufeff', ''),  # BOM
        ]
        
        for pattern, replacement in special_chars:
            if re.search(pattern, text):
                text = re.sub(pattern, replacement, text)
                corrections.append({"type": "special_char", "pattern": pattern})
        
        return text.strip(), corrections
    
    def _validate_quality(self, text: str) -> QualityDetails:
        """품질 검증 (문서 유형별 유연한 평가)"""
        if not text or len(text.strip()) < 10:
            return QualityDetails(overall_score=0.0)
        
        # 한글 비율
        text_no_space = re.sub(r'\s', '', text)
        korean_chars = len(self.KOREAN_PATTERN.findall(text))
        korean_ratio = korean_chars / len(text_no_space) if text_no_space else 0.0
        
        # 숫자/영문 비율 확인 (통계/기술 문서 판별)
        digit_count = sum(1 for c in text_no_space if c.isdigit())
        english_count = sum(1 for c in text_no_space if c.isascii() and c.isalpha())
        digit_ratio = digit_count / len(text_no_space) if text_no_space else 0
        english_ratio = english_count / len(text_no_space) if text_no_space else 0
        is_technical_doc = digit_ratio > 0.15 or english_ratio > 0.20
        
        # 한글 비율 점수화 (문서 유형 고려)
        if is_technical_doc:
            if korean_ratio >= 0.20:
                korean_score = 1.0
            elif korean_ratio >= 0.10:
                korean_score = 0.8
            elif korean_ratio >= 0.05:
                korean_score = 0.6
            else:
                korean_score = 0.3
        else:
            if 0.55 <= korean_ratio <= 0.90:
                korean_score = 1.0
            elif 0.40 <= korean_ratio < 0.55 or 0.90 < korean_ratio <= 0.95:
                korean_score = 0.7
            elif 0.25 <= korean_ratio < 0.40:
                korean_score = 0.5
            elif 0.15 <= korean_ratio < 0.25:
                korean_score = 0.3
            else:
                korean_score = 0.1
        
        # 깨진 문자 비율
        broken_chars = len(self.BROKEN_CHAR_PATTERN.findall(text))
        broken_ratio = broken_chars / len(text) if text else 0.0
        broken_score = 1.0 - min(broken_ratio * 5, 1.0)
        
        # 문장 구조 (목록/표 형식 지원)
        lines = text.split('\n')
        words = text.split()
        
        # 마크다운 표 확인
        table_lines = sum(1 for line in lines if '|' in line and line.count('|') >= 2)
        if table_lines >= 3:
            sentence_score = 1.0 if table_lines / max(len(lines), 1) >= 0.2 else 0.9
        else:
            # 목록 형식 확인
            list_pattern = re.compile(r'^\s*[-•○※ㅇ□▶▷◆◇★☆►◎⦁⦾]\s*\S')
            numbered_pattern = re.compile(r'^\s*(\d+[.)]\s*|\([가-힣\d]+\)\s*|\d+\.\d+\s*)')
            list_lines = sum(1 for line in lines if list_pattern.match(line) or numbered_pattern.match(line))
            
            if list_lines >= 3 and list_lines / max(len(lines), 1) >= 0.2:
                sentence_score = 1.0
            else:
                # 일반 문장 구조
                sentences = re.findall(r'[.!?。]\s', text)
                if sentences and len(words) >= 10:
                    avg_words = len(words) / len(sentences)
                    if 5 <= avg_words <= 30:
                        sentence_score = 1.0
                    elif 3 <= avg_words <= 50:
                        sentence_score = 0.7
                    else:
                        sentence_score = 0.4
                else:
                    non_empty = [l for l in lines if l.strip()]
                    if len(non_empty) >= 5:
                        avg_len = sum(len(l.strip()) for l in non_empty) / len(non_empty)
                        sentence_score = 0.7 if 10 <= avg_len <= 200 else 0.5
                    else:
                        sentence_score = 0.5
        
        # 반복 패턴 점수 (DOCX는 반복 패턴 문제 적음)
        repetition_score = 1.0
        
        # 인코딩 오류 검사
        has_encoding_error = detect_utf16_misencoding(text)
        encoding_score = 0.0 if has_encoding_error else 1.0
        
        # 종합 점수
        overall = (
            korean_score * 0.30 +
            broken_score * 0.25 +
            sentence_score * 0.15 +
            repetition_score * 0.10 +
            encoding_score * 0.20
        )
        
        return QualityDetails(
            korean_ratio=korean_ratio,
            broken_char_ratio=broken_ratio,
            sentence_structure=sentence_score,
            repetition_anomaly=0.0,
            encoding_error=has_encoding_error,
            korean_score=round(korean_score, 4),
            broken_score=round(broken_score, 4),
            sentence_score=round(sentence_score, 4),
            repetition_score=round(repetition_score, 4),
            encoding_score=round(encoding_score, 4),
            overall_score=round(overall, 4)
        )
    
    def _count_tokens(self, text: str) -> int:
        """토큰 수 계산"""
        try:
            import tiktoken
            enc = tiktoken.get_encoding("cl100k_base")
            return len(enc.encode(text))
        except ImportError:
            return len(text) // 2
